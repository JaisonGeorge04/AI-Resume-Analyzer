import io
import os
# pyrefly: ignore [missing-import]
import pdfplumber
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes. Tries pdfplumber first, falls back to pypdf.
    """
    text = ""
    # Try pdfplumber first
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}. Falling back to pypdf.")
        text = ""

    # Fallback to pypdf if text extraction was empty or failed
    if not text.strip():
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"pypdf fallback failed: {e}")
            raise ValueError(f"Could not parse PDF file: {e}")

    if not text.strip():
        raise ValueError("The PDF file appears to be empty or contains non-extractable text.")
        
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)
        if not text.strip():
            raise ValueError("The document appears to be empty.")
        return text.strip()
    except Exception as e:
        print(f"python-docx failed: {e}")
        raise ValueError(f"Could not parse DOCX file: {e}")

def parse_resume(filename: str, file_bytes: bytes) -> str:
    """
    Main entry point to parse resumes. Supports .pdf and .docx extensions.
    """
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")
